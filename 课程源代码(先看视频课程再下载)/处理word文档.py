# import docx


'''安装python-docx
    注意：1、库名为python-docx，并不是docx
          2、导入python-docx库只需要import docx，并不是需要import python-docx

python-docx官方手册：
    https://python-docx.readthedocs.io
'''
# f = open('实验3.docx','r',encoding='utf-8')
# print(f)
# print(f.readlines())#不能用open来打开带格式的文档

# doc = docx.Document('实验3.docx') #获取文档

# print(dir(doc))         #打印doc所有的可操作函数方法
# print(doc.paragraphs)   #打印文档中段落对象
# print(len(doc.paragraphs)) #该段文本中段落数量
# print(doc.paragraphs[0].text)  #该段中文本的字符串内容(没有样式信息)
# print(doc.paragraphs[1].text)
# print(doc.paragraphs[0].runs)  #获取对应段落中的runs个数
# print(doc.paragraphs[0].runs[0].text)
# print(doc.paragraphs[0].runs[1].text)
# print(doc.paragraphs[0].runs[2].text)
# print(doc.paragraphs[0].runs[3].text)



'''设置字符样式
add_break(break_type=WD_BREAK.LINE)
    参数：WD_BREAK.LINE:在run的后面添加一个换行
          WD_BREAK.PAGE:在run的后面添加一个换页
'''
# from docx.enum.text import WD_BREAK
# doc.paragraphs[0].runs[0].add_break(break_type=WD_BREAK.PAGE)

'''添加图片add_picture(image_path_or_stream, width=None, height=None
    参数:image_path_or_stream图片路径
    参数:width、height图片大小，指定一边，则另一边按比例缩放   
'''
# doc.paragraphs[0].runs[0].add_picture('lufei.png',width=1000000)


'''添加一个tab '''
# doc.paragraphs[0].runs[0].add_tab()


'''添加文字add_text'''
# doc.paragraphs[0].runs[2].add_text('新添加一部分内容')


'''当前run加粗bold'''
# doc.paragraphs[0].runs[0].bold = True

'''清除内容clear'''
# doc.paragraphs[0].runs[0].clear()



'''font:all_caps(字母大写)、bold(加粗)、color(字体颜色)、
    complex_script(复杂脚本)、double_strike(双删除线)、
    emboss(浮雕)、hidden(隐藏)、italic(斜体)、outline(外框)、
    size(字体大小)、name(字体名称)....更多内容请查看Font.py文件
    '''
# doc.paragraphs[0].runs[0].font.all_caps = True

# doc.paragraphs[0].runs[0].font.bold = True

# from  docx.shared import RGBColor
# doc.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 205, 205)

# doc.paragraphs[0].runs[0].font.complex_script  =True
#
# doc.paragraphs[0].runs[0].font.double_strike = True

# doc.paragraphs[0].runs[0].font.emboss = True

# doc.paragraphs[0].runs[0].font.hidden = False

# doc.paragraphs[0].runs[0].font.italic = True

# doc.paragraphs[0].runs[0].font.outline  =True

# doc.paragraphs[0].runs[0].font.size = 211124

# print(doc.paragraphs[0].runs[0].font.size.pt)

# from  docx.shared import Pt
# doc.paragraphs[0].runs[0].font.size = Pt(18)

# from docx.oxml.ns import qn
# doc.paragraphs[0].runs[0].font.name = '黑体'
# doc.paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
#


'''斜体italic'''
# doc.paragraphs[0].runs[0].italic = True
#

'''读取/写入run中的文本text'''
# print('读取的内容为：',doc.paragraphs[0].runs[0].text)
# doc.paragraphs[0].runs[0].text = '办公自动化设计'
# doc.save('实验3.docx')

'''下划线underline'''
# doc.paragraphs[0].runs[0].underline = True
# doc.save('实验3.docx')



'''项目：将实验3.docx中将不同样式的内容进行设置，
        全文英文字母字体为Times New Roman,中文为宋体，四号字体
        当样式内容小于10，则设置为红色字体，下划线
        当样式内容大于10，则设置为黑色字体，斜体
        
        实现任务的第一步：打开docx----python-docx，docx.Document()
        第二步：for  遍历doc文档段落，len()
        第三步：英文字母字体为Times New Roman，中文为宋体，四号
'''
# import docx
# from docx.oxml.ns import qn
# from  docx.shared import RGBColor
# from  docx.shared import Pt
#
# doc = docx.Document('实验3.docx') #获取文档
# for i in range(len(doc.paragraphs)):
#     for j in range(len(doc.paragraphs[i].runs)):
#         doc.paragraphs[i].runs[j].font.name = 'Times New Roman'  #英文的字体
#         doc.paragraphs[i].runs[j]._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
#         doc.paragraphs[i].runs[j].font.size = Pt(14)
#         if len(doc.paragraphs[i].runs[j].text) < 10:
#             doc.paragraphs[i].runs[j].font.color.rgb = RGBColor(255, 0, 0)
#             doc.paragraphs[i].runs[j].underline = True
#         else:
#             doc.paragraphs[i].runs[j].font.color.rgb = RGBColor(0, 0 ,0)
#             doc.paragraphs[i].runs[j].italic = True
# doc.save('实验3.docx')


'''python-docx官方手册：
    https://python-docx.readthedocs.io'''

# import docx

'''##################段落paragraphs的方法######################'''
# doc = docx.Document('实验3.docx') #获取文档

'''add_run添加run'''
# print(len(doc.paragraphs))
# doc.paragraphs[0].add_run(text='生命的追求在于不断的突破')
# doc.save('实验3.docx')


'''alignment设置段落对齐方式
LEFT(左对齐),CENTER(居中),RIGHT(右对齐),JUSTIFY(两端对齐)
DISTRIBUTE(分散对齐)
'''
# from docx.enum.text import WD_ALIGN_PARAGRAPH   #
# doc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
# doc.paragraphs[1].alignment = WD_ALIGN_PARAGRAPH.CENTER
# doc.paragraphs[2].alignment = WD_ALIGN_PARAGRAPH.RIGHT
# doc.paragraphs[3].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
# doc.paragraphs[4].alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE

#

'''clear清除段落的内容，但样式会保存'''
# doc.paragraphs[0].clear()


'''insert_paragraph_before:在之前插入一段，
    插入的内容为一个run，样式相同'''
# doc.paragraphs[0].insert_paragraph_before(text='今天是个好日子')


'''提供访问此段落格式属性（如行距和缩进）的|段落格式|对象。'''
# from docx.enum.text import WD_ALIGN_PARAGRAPH
# doc.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
#
from  docx.shared import Pt
# doc.paragraphs[0].paragraph_format.first_line_indent = Pt(70) #首行缩进
# doc.paragraphs[0].paragraph_format.left_indent = Pt(70)  #左侧缩进
# doc.paragraphs[0].paragraph_format.line_spacing = Pt(70)   #行间距

# from docx.enum.text import WD_LINE_SPACING    #按照规则设置行间距
# doc.paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE  #1.5倍行距
# doc.paragraphs[2].paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE      #2倍行距
# doc.paragraphs[5].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE      #单倍行距
# doc.paragraphs[3].paragraph_format.page_break_before = True #是否对应段落出现在下一页顶部，等效为段前分页
# doc.paragraphs[3].paragraph_format.right_indent = Pt(70) #右侧缩进
# doc.paragraphs[3].paragraph_format.space_after =  Pt(70)     #段后间距
# doc.paragraphs[3].paragraph_format.space_before=  Pt(40)     #段前间距

'''单位大小设置Inches(英寸)、Cm(厘米)、Mm(毫米)、Pt(磅)、Twips、Emu
    1In = 2.54cm
    1emu = 1/914400 inches = 1/360000cm
    1pt=0.376毫米  
    1twip = 1/20 pt
'''
# c,Mm,Pt,Twips
# doc.paragraphs[0].paragraph_format.line_spacing = Inches(1)
# doc.paragraphs[1].paragraph_format.line_spacing = Cm(1)
# doc.paragraphs[3].paragraph_format.line_spacing = Mm(1)
# doc.paragraphs[4].paragraph_format.line_spacing = Pt(1)
# doc.paragraphs[5].paragraph_format.line_spacing = Twips(1)
# doc.save('实验3.docx')



'''runs对象列表'''
# print(doc.paragraphs[0].runs)

'''text获取段落内容 | 修改段落内容'''
# print(doc.paragraphs[0].text)
# doc.paragraphs[0].text = '今天是个好日子'
# doc.save('实验3.docx')
'''在当前段落前插入一个新的段落'''
# doc.paragraphs[0]._insert_paragraph_before()
# doc.save('实验3.docx')








'''#####################样式##########################'''
''' 样式的定义：给纯文本添加颜色，字体，字号，等等多种效果的集合，
                可以提前定义好样式，独立于文本之外的对象。
    样式包括：字符样式、段落样式、表格样式、列表样式
    设置样式的好处：快速给多个内容设置相同的样式，
    
'''
'''BaseStyle'''
# import docx
# doc = docx.Document('实验3.docx') #获取文档
# print(doc.paragraphs[0].style.builtin)   #builtin用来判断是否为内置样式
# doc.paragraphs[0].style.delete()  #删除样式，同时取消对应关联的段落样式
# doc.paragraphs[0].style.hidden = False    #隐藏样式
# print(doc.paragraphs[0].style.name)       #显示样式名称
# doc.paragraphs[0].style.name = '自创样式' #改变样式名称
# print(doc.paragraphs[2].style.priority)   #段落样式优先级，在界面中显示的顺序，从小到大显示
# doc.paragraphs[2].style.priority = 20
# doc.paragraphs[0].style.quick_style = False #True：样式显示在样式库中，Flase：样式隐藏在样式库中
# print(doc.paragraphs[0].style.style_id)   #每个样式都有自己独一无二的键值
# print(doc.paragraphs[1].style.style_id)
# print(doc.paragraphs[2].style.style_id)
# print(doc.paragraphs[3].style.style_id)
# print(doc.styles)  #对象---包含很多内容
# for style_one in doc.styles:
#     print(style_one.name)
# print(doc.paragraphs[3].style.type)


'''段落样式：段落样式继承字符样式和基础样式'''
import docx
from  docx.shared import Cm,Pt

doc = docx.Document('实验3.docx')                                   #获取文档

doc.paragraphs[0].runs[0].text = '       '
doc.paragraphs[0].runs[0].underline = True
# doc.paragraphs[0].style.paragraph_format.first_line_indent = Cm(5)  #段落设置
# from docx.enum.text import WD_LINE_SPACING
# doc.paragraphs[0].style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
# doc.paragraphs[0].style.font.size = Pt(12)                          #字体设置
# from docx.oxml.ns import qn
# doc.paragraphs[0].style.font.name = 'Times New Roman'
# doc.paragraphs[0].style.font._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
#
doc.save('实验3.docx')

'''注:以上对样式的操作是针对具体的段落样式，而不是直接对样式对象操作
    document类中包含了直接对样式的操作；
'''

# import docx
#
# doc = docx.Document('实验3.docx')                                   #获取文档
# print(doc.styles.__len__())
#
# for style_one in doc.styles:
#     print(style_one.name)

'''添加样式add_style(name, style_type, builtin=False)
        参数name:添加的样式名称，数据类型为字符串
        参数：style_type，确定添加的样式是哪种样式(CHARACTER、PARAGRAPH、TABLE、LIST)
        参数:builtin,默认为False，表示为非内置样式，如果为True表示创建一个内置样式
'''
# from docx.enum.style import WD_STYLE_TYPE
# new_style = doc.styles.add_style('自创样式4',style_type = WD_STYLE_TYPE.PARAGRAPH)
# print(new_style.name)  #new_style具备段落样式所有的功能，
# print(new_style.type)
# from  docx.shared import Cm,Pt
# new_style.font.size = Pt(12)
# new_style.paragraph_format.first_line_indent = Cm(1)
# doc.save('实验3.docx')
# 注：样式中基础设定继承于段落的Normal
# 对于不同的样式，例如表格样式，默认的样式又是什么呢？

'''default(style_type)获取对应样式类型的默认样式
    参数style_type：确定样式类型(CHARACTER、PARAGRAPH、TABLE、LIST)
'''
# from docx.enum.style import WD_STYLE_TYPE
# def_style1 = doc.styles.default(WD_STYLE_TYPE.CHARACTER)
# def_style2 = doc.styles.default(WD_STYLE_TYPE.PARAGRAPH)
# def_style3 = doc.styles.default(WD_STYLE_TYPE.TABLE)
# def_style4 = doc.styles.default(WD_STYLE_TYPE.LIST)
# print(def_style1.name)
# print(def_style2.name)
# print(def_style3.name)
# print(def_style4.name)

'''get_by_id(style_id, style_type)获取对应id样式对象
    参数style_id：对应样式的id，如果没有给与id，则返回对应类型的默认样式
    参数style_type：对应样式的类型
'''
# from docx.enum.style import WD_STYLE_TYPE
# print(doc.paragraphs[0].style.style_id)  #
# style_1 = doc.styles.get_by_id('a3',WD_STYLE_TYPE.PARAGRAPH)
# print(style_1.name)


'''给段落赋予样式'''
# import docx
from docx.enum.style import WD_STYLE_TYPE

# doc = docx.Document('实验3.docx')
# doc.paragraphs[0].style = doc.styles.get_by_id('a3',WD_STYLE_TYPE.PARAGRAPH) #获取到对应ID的样式对象
# doc.paragraphs[0].style = '自创样式3'

# for style_1  in doc.styles:
#     print(style_1.name)
# for i in range(len(doc.paragraphs)):
#     doc.paragraphs[i].style = doc.styles.get_by_id('a3', WD_STYLE_TYPE.PARAGRAPH)
# doc.save('实验3.docx')


'''############Document####################'''
# import docx

# doc = docx.Document('实验3.docx')

'''add_heading(text="", level=1)在文档的最后添加一个标题内容
    参数：text:需要添加的内容文本
    参数：level:使用哪种标题样式，注：这里只能使用已存在的标题
'''
# doc.add_heading(text='python学习',level= 1)
# doc.save('实验3.docx')
# for style in doc.styles:
#     print(style.name)

'''add_paragraph(text='', style=None)'''
# s= '办公自动化大大提高办公效率'
# parg = doc.add_paragraph(text=s, style ='Heading 1')
# doc.save('实验3.docx')


'''add_picture(image_path_or_stream, width=None, height=None)
    作用：在文档的末尾段落插入一个图片，单独为一段；如果未指定宽度和高度，
    则图片将以其原始大小显示。如果只指定了一个，它将用于计算比例因子，
    然后应用于未指定的维度，保持图像的纵横比。
    参数image_path_or_stream：图片地址
    参数width：图片的宽度
    参数height:图片的高度
'''
from docx.shared import Cm
# from docx.enum.text import WD_ALIGN_PARAGRAPH
# doc.add_picture('lufei.png',width=Cm(5))
# doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
# doc.save('实验3.docx')


'''add_section分节符:是指为表示节的结尾插入的标记。
        作用就是将文章分成不同的“节”。
        插入一个分节符表示一个“节”在此结束，分节符之后将开始新的一节。
        分节符包含节的格式设置元素，如页边距、页面的方向、页眉和页脚，以及页码的顺序。
        分节符用一条横贯屏幕的虚双线表示。
    add_section(start_type=WD_SECTION.NEW_PAGE)
        参数start_type:插入节的位置，参数值是docx/enum/WD_SECTION中的成员
            CONTINUOUS：连续，
            NEW_COLUMN：新列，
            NEW_PAGE：在document最后添加一个节
            EVEN_PAGE：偶数页
            ODD_PAGE：奇数页
'''
from docx.enum.section import WD_SECTION
# section_1 = doc.add_section(start_type = WD_SECTION.NEW_PAGE)
# print(section_1.start_type)  #section_1是一个节对象，具备多种方法
# print(doc.sections)
# doc.save('实验3.docx')

'''sections:获取docx文档中的节对象'''
# for section_one in doc.sections:
#     print(section_one.start_type)

'''section节：一个document至少包含一个节，'''
# from  docx.shared import Cm
'''bottom_margin:设置节内的所有页底部边距'''
# doc.sections[0].bottom_margin = Cm(0.5)
# doc.save('实验3.docx')


'''section节中创建页眉(header)页脚(footer)
header()继承_Header继承_BaseHeaderFooter继承BlockItemContainer
footer()继承_Footer继承_BaseHeaderFooter继承BlockItemContainer
因此header()和footer()创建后可以使用段落来进行设置
'''
# doc.sections[0].header.paragraphs[0].text= '这是第一个页眉!!!'
# from docx.enum.text import WD_ALIGN_PARAGRAPH
# doc.sections[0].header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# doc.sections[1].header.is_linked_to_previous = False   #是否链接到上一个节
# doc.sections[1].header.paragraphs[0].text= '这是第二个页眉'
# doc.sections[0].footer.paragraphs[0].text = '创建了一个页脚'
# doc.save('实验3.docx')

'''add_table(rows, cols, style=None)在document最后添加一个表
    参数rows：添加的表行数
    参数cols：添加的表列数
    参数style:可以是段落样式对象或段落样式名称,如果为None,则表将继承文档的默认表样式。
'''
# new_table = doc.add_table(3,4)
# print(dir(new_table.cell(1,2)))
# new_table.cell(1,2).text = 'nihao '
# print(doc.tables)
# doc.save('实验3.docx')

'''##################实战项目#####################
北京奔驰汽车公司需要给用户王强发送一份电子word版本的使用说明书！
格式说明：
    其中封面包含标题1汽车说明书大全和一个表格，表格中填写王强的信息，
    标题1的格式：居中对齐、字体36磅、加粗、中文宋体、西文Times New Roman
    表格内容：第一列内容右对齐，第二列左对齐，字体18磅、加粗、中文宋体、西文Times New Roman
    篇一篇二篇三开头为标题2，左对齐、字体18磅、加粗、中文宋体、西文Times New Roman
    正文两端对齐、字体18磅、加粗、中文宋体、西文Times New Roman
    以标题二的内容加入页眉，并且首页不加入页眉，各自节中显示自己的对应标题内容
'''
# import docx
# from docx.enum.style import WD_STYLE_TYPE
# from  docx.shared import Pt
# from docx.oxml.ns import qn
# from docx.enum.text import WD_ALIGN_PARAGRAPH
#
# def style_design(name,alig,size,bold = True):
#     table_style = doc_user.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
#     table_style.paragraph_format.alignment = alig
#     table_style.font.size = Pt(size)
#     table_style.font.bold = bold
#     table_style.font.name = 'Times New Roman'
#     table_style.font._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
#
# def fengmian(doc_user):
#     '''封面设计'''
#     style_design('标题1',WD_ALIGN_PARAGRAPH.CENTER,36)
#     doc_user.add_paragraph(text='汽车说明书大全', style='标题1')
#     doc_user.add_paragraph('\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n')
#     user_info = doc_user.add_table(2, 2)
#     style_design('表格样式1', WD_ALIGN_PARAGRAPH.RIGHT, 18)
#     style_design('表格样式2', WD_ALIGN_PARAGRAPH.LEFT, 18)
#     user_info.cell(0, 0).add_paragraph('汽车车主：', style='表格样式1')
#     user_info.cell(0, 1).add_paragraph('王强先生', style='表格样式2')
#     user_info.cell(1, 0).add_paragraph('汽车车型：', style='表格样式1')
#     user_info.cell(1, 1).add_paragraph('奔驰C200', style='表格样式2')
#
#
# def zhengwen(doc_template,doc_user):
#     style_design('标题2',WD_ALIGN_PARAGRAPH.LEFT,18)
#     style_design('正文1', WD_ALIGN_PARAGRAPH.JUSTIFY,12,bold =False)
#     for parg in doc_template.paragraphs[8:]:
#         if '篇' in parg.text:
#             new_section = doc_user.add_section()
#             new_section.header.is_linked_to_previous = False
#             new_section.header.paragraphs[0].text = parg.text
#             new_section.header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
#             doc_user.add_paragraph(parg.text,'标题2')
#             continue
#         doc_user.add_paragraph(parg.text,'正文1')
#
#
# doc_template = docx.Document('word项目-汽车用户手册\北京奔驰用户手册模板.docx')
# doc_user = docx.Document()   #创建一个新的doc文档，
# fengmian(doc_user)
# zhengwen(doc_template,doc_user)
# doc_user.save('word项目-汽车用户手册\北京奔驰用户手册-王强先生.docx')

'''实战二
在实战一的基础上将excel表格中所有客户信息都制作成word文档,
文档名：为北京奔驰用户手册-客户名先生-汽车车型.docx
'''
# import openpyxl
# import docx
# from docx.enum.style import WD_STYLE_TYPE
# from  docx.shared import Pt
# from docx.oxml.ns import qn
# from docx.enum.text import WD_ALIGN_PARAGRAPH
#
# def style_design(name,alig,size,bold = True):
#     table_style = doc_user.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
#     table_style.paragraph_format.alignment = alig
#     table_style.font.size = Pt(size)
#     table_style.font.bold = bold
#     table_style.font.name = 'Times New Roman'
#     table_style.font._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
#
# def fengmian(doc_user,chezhu,chexing):
#     '''封面设计'''
#     style_design('标题1',WD_ALIGN_PARAGRAPH.CENTER,36)
#     doc_user.add_paragraph(text='汽车说明书大全', style='标题1')
#     doc_user.add_paragraph('\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n')
#     user_info = doc_user.add_table(2, 2)
#     style_design('表格样式1', WD_ALIGN_PARAGRAPH.RIGHT, 18)
#     style_design('表格样式2', WD_ALIGN_PARAGRAPH.LEFT, 18)
#     user_info.cell(0, 0).add_paragraph('汽车车主：', style='表格样式1')
#     user_info.cell(0, 1).add_paragraph(chezhu, style='表格样式2')
#     user_info.cell(1, 0).add_paragraph('汽车车型：', style='表格样式1')
#     user_info.cell(1, 1).add_paragraph(chexing, style='表格样式2')
#
#
# def zhengwen(doc_template,doc_user):
#     style_design('标题2',WD_ALIGN_PARAGRAPH.LEFT,18)
#     style_design('正文1', WD_ALIGN_PARAGRAPH.JUSTIFY,12,bold =False)
#     for parg in doc_template.paragraphs[8:]:
#         if '篇' in parg.text:
#             new_section = doc_user.add_section()
#             new_section.header.is_linked_to_previous = False
#             new_section.header.paragraphs[0].text = parg.text
#             new_section.header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
#             doc_user.add_paragraph(parg.text,'标题2')
#             continue
#         doc_user.add_paragraph(parg.text,'正文1')
#
#
# doc_template = docx.Document('word项目-汽车用户手册\北京奔驰用户手册模板.docx')
# ex_file = openpyxl.load_workbook('word项目-汽车用户手册\客户信息.xlsx')
# sheet_1 = ex_file.get_sheet_by_name('Sheet1')
# for line in sheet_1['A1':'B40']:
#     doc_user = docx.Document()
#     fengmian(doc_user,line[0].value,line[1].value)
#     zhengwen(doc_template,doc_user)
#     info = 'word项目-汽车用户手册\北京奔驰用户手册-' + line[0].value+'先生-'+line[1].value+'.docx'
#     doc_user.save(info)